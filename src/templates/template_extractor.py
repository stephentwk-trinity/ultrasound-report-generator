import json
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from .template_model import Template


class TemplateExtractor:
    def __init__(self):
        self.templates: List[Template] = []

    def extract_from_docx(self, docx_path: Path) -> List[Template]:
        """Extract templates from the Word document with advanced parsing."""
        document = Document(docx_path)
        self.templates = []
        
        # Get all body elements (paragraphs and tables)
        body_elements = self._get_body_elements(document)
        
        template_id = 1
        i = 0
        seen_templates = set()  # Track unique templates to avoid duplicates
        
        while i < len(body_elements):
            element = body_elements[i]
            
            # Check if this is a paragraph with "TEMPLATE ONLY" marker
            if isinstance(element, Paragraph) and "TEMPLATE ONLY" in element.text.upper():
                # Extract variation name from parentheses
                variation = self._extract_variation(element.text)
                
                # Parse all templates under this "TEMPLATE ONLY" section
                templates_data, next_idx = self._parse_template_section(
                    body_elements, i + 1, variation
                )
                
                # Create Template objects (with deduplication)
                for template_data in templates_data:
                    # Create a unique key for this template
                    template_key = (
                        template_data['name'],
                        variation,
                        tuple(sorted(template_data['sections'].items()))
                    )
                    
                    # Only add if we haven't seen this exact template before
                    if template_key not in seen_templates:
                        template = Template(
                            id=f"template_{template_id:03d}",
                            name=template_data['name'],
                            sections=template_data['sections'],
                            metadata={
                                "source": "US Report templates.docx",
                                "variation": variation,
                                "extracted_at": "2025-10-28"
                            }
                        )
                        self.templates.append(template)
                        seen_templates.add(template_key)
                        template_id += 1
                i = next_idx
            else:
                i += 1
        
        return self.templates

    def _get_body_elements(self, document: Document) -> List:
        """Get all body elements (paragraphs and tables) in order."""
        body_elements = []
        
        # Access the document's body element
        for element in document.element.body:
            if element.tag.endswith('p'):  # Paragraph
                # Find the corresponding Paragraph object
                for para in document.paragraphs:
                    if para._element == element:
                        body_elements.append(para)
                        break
            elif element.tag.endswith('tbl'):  # Table
                # Find the corresponding Table object
                for table in document.tables:
                    if table._element == element:
                        body_elements.append(table)
                        break
        
        return body_elements

    def _extract_variation(self, text: str) -> str:
        """Extract variation name from parentheses in TEMPLATE ONLY line."""
        match = re.search(r'\(([^)]+)\)', text)
        if match:
            return match.group(1).strip()
        return ""

    def _parse_template_section(self, body_elements: List, start_idx: int, 
                                variation: str) -> tuple[List[Dict], int]:
        """
        Parse all template examples under a single TEMPLATE ONLY section.
        Returns list of template data dicts and the next index to process.
        """
        templates_data = []
        current_template = None
        current_section = ""
        current_content = []
        main_heading = ""  # Track the main heading to detect new examples
        
        i = start_idx
        while i < len(body_elements):
            element = body_elements[i]
            
            # Check if we've reached the next "TEMPLATE ONLY" marker
            if isinstance(element, Paragraph) and "TEMPLATE ONLY" in element.text.upper():
                # Save current template before breaking
                if current_template and current_section:
                    current_template['sections'][current_section] = "\n".join(current_content).strip()
                if current_template:
                    templates_data.append(current_template)
                break
            
            # Handle paragraphs
            if isinstance(element, Paragraph):
                text = element.text.strip()
                
                if not text:
                    i += 1
                    continue
                
                # Define known section headers
                known_sections = ["CLINICAL INFORMATION", "FINDINGS", "COMMENTS", "IMPRESSION", "PROCEDURE", "SUMMARY", "L.M.P."]
                text_without_colon = text.rstrip(':')
                
                # Check if this is a known section header
                if text_without_colon in known_sections or (text.isupper() and text.endswith(':')):
                    if current_template:
                        # Save previous section
                        if current_section and current_content:
                            current_template['sections'][current_section] = "\n".join(current_content).strip()
                        
                        # Start new section
                        current_section = text_without_colon
                        current_content = []
                
                # Check if this could be a main heading (all caps, substantial text, NOT a known section)
                elif text.isupper() and len(text) > 10 and text_without_colon not in known_sections:
                    # If we already have a main heading and this matches it, start a new template
                    if main_heading and text == main_heading and current_template:
                        # Save current section
                        if current_section and current_content:
                            current_template['sections'][current_section] = "\n".join(current_content).strip()
                        
                        # Save current template
                        if current_template:
                            templates_data.append(current_template)
                        
                        # Start new template with same variation
                        current_template = {
                            'name': text,
                            'sections': {}
                        }
                        current_section = ""
                        current_content = []
                    elif not main_heading:
                        # This is the first main heading - start first template
                        main_heading = text
                        current_template = {
                            'name': text,
                            'sections': {}
                        }
                        current_section = ""
                        current_content = []
                # Regular content
                elif current_section and text:
                    current_content.append(text)
            
            # Handle tables
            elif isinstance(element, Table):
                if current_template and current_section:
                    # Convert table to markdown and add to current section
                    markdown_table = self._table_to_markdown(element)
                    current_content.append(markdown_table)
            
            i += 1
        
        # Save last section and template
        if current_template and current_section and current_content:
            current_template['sections'][current_section] = "\n".join(current_content).strip()
        if current_template:
            templates_data.append(current_template)
        
        return templates_data, i

    def _table_to_markdown(self, table: Table) -> str:
        """Convert a Word table to markdown format."""
        markdown_lines = []
        
        # Process header row (first row)
        if len(table.rows) > 0:
            header_cells = [cell.text.strip() for cell in table.rows[0].cells]
            markdown_lines.append("| " + " | ".join(header_cells) + " |")
            markdown_lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
        
        # Process data rows
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            markdown_lines.append("| " + " | ".join(cells) + " |")
        
        return "\n" + "\n".join(markdown_lines) + "\n"

    def save_to_json(self, templates: List[Template], output_path: Path):
        """Serialize templates to JSON file."""
        # Convert templates to dictionaries
        templates_data = []
        for template in templates:
            template_dict = {
                "id": template.id,
                "name": template.name,
                "sections": template.sections,
                "metadata": template.metadata
            }
            templates_data.append(template_dict)
        
        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Save to JSON
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(templates_data, f, indent=2, ensure_ascii=False)